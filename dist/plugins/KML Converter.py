import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json
import os
import re
from pathlib import Path
import threading
from datetime import datetime
import logging
import math

# Проверяем наличие необходимых библиотек
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.error("PyPDF2 не установлен. Установите: pip install PyPDF2")

try:
    import simplekml
    KML_SUPPORT = True
except ImportError:
    KML_SUPPORT = False
    logging.error("simplekml не установлен. Установите: pip install simplekml")

# Добавляем поддержку DOC/DOCX
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.error("python-docx не установлен. Установите: pip install python-docx")

class KMLCreatorPlugin:
    def __init__(self, settings, root):
        self.settings = settings
        self.root = root
        self.kml_data = None
        self.processed_data = []
        self.loaded_files = []
        
        # Настройки по умолчанию
        self.coord_format = "standard"  # standard или geographic
        self.color_settings = {
            "circle_color": "blue",
            "polygon_color": "yellow"
        }
        
        # Загрузка сохраненных настроек
        self.load_settings()
        
        # Список доступных цветов с названиями
        self.available_colors = {
            "Красный": "red",
            "Зеленый": "green",
            "Синий": "blue", 
            "Желтый": "yellow",
            "Оранжевый": "orange",
            "Фиолетовый": "purple",
            "Розовый": "pink",
            "Коричневый": "brown",
            "Черный": "black",
            "Белый": "white",
            "Голубой": "cyan",
            "Лаймовый": "lime"
        }
    
    def get_tab_name(self):
        return "KML Converter 1.4"
    
    def create_tab(self):
        # Создаем основной контейнер с разделителем
        main_container = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Левая панель - настройки
        left_frame = ttk.Frame(main_container)
        main_container.add(left_frame, weight=1)
        
        # Правая панель - обработка и результаты
        right_frame = ttk.Frame(main_container)
        main_container.add(right_frame, weight=1)
        
        # Устанавливаем начальное соотношение (60% настройки, 40% результаты)
        main_container.sashpos(0, int(self.root.winfo_width() * 0.6))
        
        # Создаем интерфейс левой и правой панелей
        self.create_left_interface(left_frame)
        self.create_right_interface(right_frame)
        
        return main_container
    
    def create_left_interface(self, parent):
        """Создание левой панели с настройками"""
        # Создаем скроллируемую область для левой панели
        left_canvas = tk.Canvas(parent)
        left_scrollbar = ttk.Scrollbar(parent, orient="vertical", command=left_canvas.yview)
        scrollable_left_frame = ttk.Frame(left_canvas)
        
        scrollable_left_frame.bind(
            "<Configure>",
            lambda e: left_canvas.configure(scrollregion=left_canvas.bbox("all"))
        )
        
        left_canvas.create_window((0, 0), window=scrollable_left_frame, anchor="nw")
        left_canvas.configure(yscrollcommand=left_scrollbar.set)
        
        left_canvas.pack(side="left", fill="both", expand=True)
        left_scrollbar.pack(side="right", fill="y")
        
        # Заголовок
        title_label = ttk.Label(scrollable_left_frame, text="KML Converter - Создание KML файлов", 
                               font=('Arial', 12, 'bold'))
        title_label.pack(anchor=tk.W, pady=(0, 10))
        
        # Проверка зависимостей
        if not KML_SUPPORT:
            warning_frame = ttk.Frame(scrollable_left_frame)
            warning_frame.pack(fill=tk.X, pady=(0, 10))
            
            warning_text = "ВНИМАНИЕ: simplekml не установлен!\nУстановите: pip install simplekml"
            
            ttk.Label(warning_frame, text=warning_text, foreground="red", 
                     justify=tk.LEFT, wraplength=600).pack(anchor=tk.W)
        
        # Фрейм настроек формата координат
        format_frame = ttk.LabelFrame(scrollable_left_frame, text="Настройки формата координат")
        format_frame.pack(fill=tk.X, pady=(0, 10))
        
        format_buttons = ttk.Frame(format_frame)
        format_buttons.pack(fill=tk.X, padx=5, pady=5)
        
        self.format_var = tk.StringVar(value=self.coord_format)
        
        ttk.Radiobutton(format_buttons, text="Стандартный (564144N0523236E)", 
                       variable=self.format_var, value="standard",
                       command=self.on_format_changed).pack(anchor=tk.W, pady=2)
        ttk.Radiobutton(format_buttons, text="Географический (56.1234, 52.5678)", 
                       variable=self.format_var, value="geographic",
                       command=self.on_format_changed).pack(anchor=tk.W, pady=2)
        
        # Подсказка о форматах
        hint_label = ttk.Label(format_frame, 
                              text="Стандартный: 564144N0523236E (градусы, минуты, секунды)\nГеографический: 56.1234, 52.5678 (десятичные градусы) или 56;331195 с. ш. ; 46;545811 в. д.",
                              font=('Arial', 9), foreground="gray", justify=tk.LEFT, wraplength=800)
        hint_label.pack(anchor=tk.W, padx=5, pady=(0, 5))
        
        # Создаем Notebook для вкладок
        notebook = ttk.Notebook(scrollable_left_frame)
        notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Вкладка для ручного ввода координат
        manual_tab = ttk.Frame(notebook)
        notebook.add(manual_tab, text="Ручной ввод")
        self.create_manual_tab(manual_tab)
        
        # Вкладка для загрузки файлов
        files_tab = ttk.Frame(notebook)
        notebook.add(files_tab, text="Из файлов")
        self.create_files_tab(files_tab)
        
        # Фрейм настроек цветов
        colors_frame = ttk.LabelFrame(scrollable_left_frame, text="Настройки цветов")
        colors_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.create_color_settings(colors_frame)
        
        # Статус бар в левой панели
        self.status_var = tk.StringVar(value="Готов к работе")
        status_bar = ttk.Label(scrollable_left_frame, textvariable=self.status_var, 
                              relief=tk.SUNKEN, style='TLabel')
        status_bar.pack(fill=tk.X, pady=(5, 0))
    
    def create_right_interface(self, parent):
        """Создание правой панели с обработкой и результатами"""
        # Фрейм обработки данных
        process_frame = ttk.LabelFrame(parent, text="Обработка данных")
        process_frame.pack(fill=tk.X, pady=5, padx=5)
        
        process_buttons = ttk.Frame(process_frame)
        process_buttons.pack(fill=tk.X, padx=5, pady=5)
        
        # Кнопки в строку с равномерным распределением
        process_buttons.columnconfigure(0, weight=1)
        process_buttons.columnconfigure(1, weight=1)
        process_buttons.columnconfigure(2, weight=1)
        
        ttk.Button(process_buttons, text="Обработать данные", 
                  command=self.process_data, state="normal" if KML_SUPPORT else "disabled").grid(row=0, column=0, padx=2, sticky="ew")
        ttk.Button(process_buttons, text="Инструкция", 
                  command=self.show_instructions).grid(row=0, column=1, padx=2, sticky="ew")
        ttk.Button(process_buttons, text="Очистить все", 
                  command=self.clear_all).grid(row=0, column=2, padx=2, sticky="ew")
        
        # Область результатов
        result_frame = ttk.LabelFrame(parent, text="Результаты обработки")
        result_frame.pack(fill=tk.BOTH, expand=True, pady=5, padx=5)
        result_frame.columnconfigure(0, weight=1)
        result_frame.rowconfigure(1, weight=1)  # Текстовое поле получает вес
        
        # Кнопка экспорта в правой панели
        export_frame = ttk.Frame(result_frame)
        export_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(export_frame, text="Экспорт KML", 
                  command=self.export_kml, state="normal" if KML_SUPPORT else "disabled").pack(side=tk.LEFT, padx=5)
        
        # Текстовое поле с прокруткой для результатов
        result_text_frame = ttk.Frame(result_frame)
        result_text_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        result_text_frame.columnconfigure(0, weight=1)
        result_text_frame.rowconfigure(0, weight=1)
        
        self.result_text = tk.Text(result_text_frame, wrap=tk.WORD, state=tk.DISABLED)
        result_scrollbar = ttk.Scrollbar(result_text_frame, orient=tk.VERTICAL, 
                                       command=self.result_text.yview)
        self.result_text.configure(yscrollcommand=result_scrollbar.set)
        
        self.result_text.grid(row=0, column=0, sticky="nsew")
        result_scrollbar.grid(row=0, column=1, sticky="ns")
    
    def on_format_changed(self):
        """Обработка изменения формата координат"""
        self.coord_format = self.format_var.get()
        self.save_settings()
        self.update_hint_label()
        self.update_status(f"Формат координат изменен на: {'Стандартный' if self.coord_format == 'standard' else 'Географический'}")
    
    def create_manual_tab(self, parent):
        """Создание вкладки для ручного ввода координат"""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)
        
        # Фрейм для ввода координат
        input_frame = ttk.LabelFrame(parent, text="Ввод координат")
        input_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        input_frame.columnconfigure(0, weight=1)
        input_frame.rowconfigure(1, weight=1)
        
        # Контейнер для текстового поля и подсказки
        content_frame = ttk.Frame(input_frame)
        content_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        content_frame.columnconfigure(0, weight=1)
        content_frame.rowconfigure(1, weight=1)
        
        # Добавляем подсказку
        self.hint_label = ttk.Label(content_frame, text="", font=('Arial', 9), 
                                   foreground="gray", justify=tk.LEFT, wraplength=800)
        self.hint_label.grid(row=0, column=0, sticky="w", pady=(0, 5))
        self.update_hint_label()
        
        # Фрейм для текстового поля и скроллбара
        text_container = ttk.Frame(content_frame)
        text_container.grid(row=1, column=0, sticky="nsew", pady=(0, 5))
        text_container.columnconfigure(0, weight=1)
        text_container.rowconfigure(0, weight=1)
        
        # Текстовое поле
        self.coord_text = tk.Text(text_container, wrap=tk.WORD, height=10)
        
        # Настраиваем прокрутку
        text_scrollbar = ttk.Scrollbar(text_container, orient=tk.VERTICAL, command=self.coord_text.yview)
        self.coord_text.configure(yscrollcommand=text_scrollbar.set)
        
        self.coord_text.grid(row=0, column=0, sticky="nsew")
        text_scrollbar.grid(row=0, column=1, sticky="ns")
        
        # Привязываем комбинации клавиш для русской и английской раскладки
        self.coord_text.bind('<Control-c>', self.copy_text)
        self.coord_text.bind('<Control-C>', self.copy_text)
        self.coord_text.bind('<Control-v>', self.paste_text)
        self.coord_text.bind('<Control-V>', self.paste_text)
        self.coord_text.bind('<Key>', self.handle_keypress)
        
        # Кнопки для ручного ввода
        manual_buttons = ttk.Frame(content_frame)
        manual_buttons.grid(row=2, column=0, sticky="ew", pady=5)
        manual_buttons.columnconfigure(0, weight=1)
        
        ttk.Button(manual_buttons, text="Очистить поле", 
                  command=self.clear_manual_text).pack(side=tk.LEFT, padx=2)
    
    def update_hint_label(self):
        """Обновление подсказки в зависимости от выбранного формата"""
        if self.coord_format == "standard":
            hint_text = "Вставьте координаты в формате: 524651N0500052E\nПоддерживаются форматы: отдельные точки, круговые зоны (ОКРУЖНОСТЬ РАДИУС 10 КМ ЦЕНТР 524651N0500052E), полигоны (РАЙОН 524651N0500052E 524751N0500152E ...)\n\nПримечание: отдельные точки не будут отображаться в KML файле."
        else:
            hint_text = "Вставьте координаты в формате: 56.1234, 52.5678 или 56;331195 с. ш. ; 46;545811 в. д.\nПоддерживаются форматы: отдельные точки, круговые зоны (ОКРУЖНОСТЬ РАДИУС 10 КМ ЦЕНТР 56.1234, 52.5678), полигоны (РАЙОН 56.1234,52.5678 56.2234,52.6678 ...)\n\nПримечание: отдельные точки не будут отображаться в KML файле."
        
        if hasattr(self, 'hint_label'):
            self.hint_label.config(text=hint_text)
    
    def create_files_tab(self, parent):
        """Создание вкладки для загрузки файлов"""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)
        
        # Проверка поддержки PDF и DOCX
        warning_shown = False
        
        if not PDF_SUPPORT:
            warning_frame = ttk.Frame(parent)
            warning_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
            warning_frame.columnconfigure(0, weight=1)
            
            warning_text = "PyPDF2 не установлен. Функция загрузки PDF недоступна.\nУстановите: pip install PyPDF2\nTXT и DOCX файлы можно загружать без установки PyPDF2."
            ttk.Label(warning_frame, text=warning_text, foreground="red", 
                     justify=tk.LEFT, wraplength=600).pack(anchor=tk.W, fill=tk.X)
            warning_shown = True
        
        if not DOCX_SUPPORT:
            warning_frame = ttk.Frame(parent)
            row_idx = 1 if warning_shown else 0
            warning_frame.grid(row=row_idx, column=0, sticky="ew", padx=5, pady=(0 if warning_shown else 5))
            warning_frame.columnconfigure(0, weight=1)
            
            warning_text = "python-docx не установлен. Функция загрузки DOC/DOCX недоступна.\nУстановите: pip install python-docx\nTXT и PDF файлы можно загружать без установки python-docx."
            ttk.Label(warning_frame, text=warning_text, foreground="red", 
                     justify=tk.LEFT, wraplength=600).pack(anchor=tk.W, fill=tk.X)
        
        # Фрейм загрузки файлов
        upload_frame = ttk.LabelFrame(parent, text="Загрузка файлов (PDF, TXT, DOC, DOCX)")
        row_idx = 2 if warning_shown and not DOCX_SUPPORT else (1 if warning_shown or not DOCX_SUPPORT else 0)
        upload_frame.grid(row=row_idx, column=0, sticky="nsew", padx=5, pady=5)
        upload_frame.columnconfigure(0, weight=1)
        upload_frame.rowconfigure(2, weight=1)
        
        # Подсказка
        hint_label = ttk.Label(upload_frame, 
                              text="Поддерживаются файлы PDF, TXT, DOC и DOCX. В файлах ищутся координаты, круговые зоны и полигоны.\nОтдельные точки не будут отображаться в KML файле.",
                              font=('Arial', 9), foreground="gray", justify=tk.LEFT, wraplength=800)
        hint_label.grid(row=0, column=0, sticky="w", padx=5, pady=(5, 0))
        
        # Кнопки загрузки
        button_frame = ttk.Frame(upload_frame)
        button_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=5)
        button_frame.columnconfigure(0, weight=1)
        
        ttk.Button(button_frame, text="Добавить файлы", 
                  command=self.add_files).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Очистить список", 
                  command=self.clear_files).pack(side=tk.LEFT, padx=2)
        
        # Список загруженных файлов
        list_frame = ttk.Frame(upload_frame)
        list_frame.grid(row=2, column=0, sticky="nsew", padx=5, pady=5)
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        
        self.files_listbox = tk.Listbox(list_frame, height=8)
        files_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.files_listbox.yview)
        self.files_listbox.configure(yscrollcommand=files_scrollbar.set)
        
        self.files_listbox.grid(row=0, column=0, sticky="nsew")
        files_scrollbar.grid(row=0, column=1, sticky="ns")
    
    def create_color_settings(self, parent):
        """Создание настроек цветов"""
        # Сетка для цветов
        colors_grid = ttk.Frame(parent)
        colors_grid.pack(fill=tk.X, padx=5, pady=5)
        colors_grid.columnconfigure(1, weight=1)
        
        # Круговые зоны
        ttk.Label(colors_grid, text="Круговые зоны:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        self.circle_color_combo = self.create_color_combobox(colors_grid, "circle_color", 0, 1)
        
        # Полигональные зоны
        ttk.Label(colors_grid, text="Полигональные зоны:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        self.polygon_color_combo = self.create_color_combobox(colors_grid, "polygon_color", 1, 1)
        
        # Кнопка сброса цветов
        ttk.Button(colors_grid, text="Сбросить цвета", 
                  command=self.reset_colors).grid(row=2, column=0, columnspan=2, pady=10, sticky="w")
    
    def create_color_combobox(self, parent, color_key, row, column):
        """Создание комбобокса для выбора цвета"""
        # Фрейм для комбобокса и предпросмотра
        color_frame = ttk.Frame(parent)
        color_frame.grid(row=row, column=column, sticky="w", padx=5, pady=2)
        
        # Предпросмотр цвета
        color_preview = tk.Frame(color_frame, width=20, height=20, 
                                background=self.color_settings.get(color_key, "green"))
        color_preview.pack(side=tk.LEFT, padx=(0, 5))
        color_preview.pack_propagate(False)
        
        # Комбобокс
        color_names = list(self.available_colors.keys())
        current_color_name = self.get_color_name(self.color_settings.get(color_key, "green"))
        
        combo = ttk.Combobox(color_frame, values=color_names, width=15, state="readonly")
        combo.set(current_color_name)
        combo.pack(side=tk.LEFT)
        
        # Привязка события изменения
        combo.bind('<<ComboboxSelected>>', 
                  lambda e, key=color_key, preview=color_preview: 
                  self.on_color_changed(key, combo.get(), preview))
        
        return combo
    
    def get_color_name(self, color_value):
        """Получение названия цвета по значению"""
        for name, value in self.available_colors.items():
            if value == color_value:
                return name
        return "Пользовательский"
    
    def on_color_changed(self, color_key, color_name, color_preview):
        """Обработка изменения цвета через комбобокс"""
        if color_name in self.available_colors:
            color_value = self.available_colors[color_name]
            self.color_settings[color_key] = color_value
            color_preview.config(background=color_value)
            self.save_settings()
    
    def reset_colors(self):
        """Сброс цветов к значениям по умолчанию"""
        default_colors = {
            "circle_color": "blue",
            "polygon_color": "yellow"
        }
        
        self.color_settings.update(default_colors)
        self.save_settings()
        
        # Обновляем комбобоксы
        self.circle_color_combo.set(self.get_color_name("blue"))
        self.polygon_color_combo.set(self.get_color_name("yellow"))
        
        # Обновляем предпросмотры
        for widget in self.circle_color_combo.master.winfo_children():
            if isinstance(widget, tk.Frame):
                widget.config(background="blue")
                break
                
        for widget in self.polygon_color_combo.master.winfo_children():
            if isinstance(widget, tk.Frame):
                widget.config(background="yellow")
                break
    
    def handle_keypress(self, event):
        """Обработка нажатий клавиш для поддержки русского Ctrl+C/Ctrl+V"""
        # Проверяем, нажаты ли Ctrl+C или Ctrl+V в русской раскладке
        if event.state & 0x4:  # Проверяем, что нажат Ctrl
            char = event.char.lower()
            if char == 'с' or char == 'c':  # Русская 'с' или английская 'c'
                self.copy_text()
                return "break"
            elif char == 'м' or char == 'v':  # Русская 'м' или английская 'v'
                self.paste_text()
                return "break"
        return None
    
    def copy_text(self, event=None):
        """Копирование текста"""
        try:
            # Получаем выделенный текст
            if self.coord_text.tag_ranges(tk.SEL):
                selected_text = self.coord_text.get(tk.SEL_FIRST, tk.SEL_LAST)
                self.root.clipboard_clear()
                self.root.clipboard_append(selected_text)
        except:
            # Если ничего не выделено
            pass
        
        # Всегда предотвращаем стандартное поведение
        if event:
            return "break"
    
    def paste_text(self, event=None):
        """Вставка текста"""
        try:
            clipboard_text = self.root.clipboard_get()
            # Если есть выделение, заменяем его
            if self.coord_text.tag_ranges(tk.SEL):
                self.coord_text.delete(tk.SEL_FIRST, tk.SEL_LAST)
            self.coord_text.insert(tk.INSERT, clipboard_text)
        except:
            pass
        
        # Всегда предотвращаем стандартное поведение
        if event:
            return "break"
    
    def clear_manual_text(self):
        """Очистка текстового поля ручного ввода"""
        self.coord_text.delete(1.0, tk.END)
        self.update_status("Текстовое поле очищено")
    
    def clear_files(self):
        """Очистка списка файлов"""
        self.loaded_files.clear()
        if hasattr(self, 'files_listbox'):
            self.files_listbox.delete(0, tk.END)
        self.update_status("Список файлов очищен")
    
    def clear_all(self):
        """Очистка всех данных"""
        self.clear_manual_text()
        self.clear_files()
        self.kml_data = None
        self.processed_data = []
        self.result_text.config(state=tk.NORMAL)
        self.result_text.delete(1.0, tk.END)
        self.result_text.config(state=tk.DISABLED)
        self.update_status("Все данные очищены")
    
    def add_files(self):
        """Добавление файлов (PDF, TXT, DOC, DOCX)"""
        filetypes = [
            ("PDF files", "*.pdf"), 
            ("TXT files", "*.txt"), 
            ("Word documents", "*.doc;*.docx"),
            ("All files", "*.*")
        ]
        
        files = filedialog.askopenfilenames(
            title="Выберите файлы (PDF, TXT, DOC или DOCX)",
            filetypes=filetypes
        )
        
        for file_path in files:
            if file_path not in self.loaded_files:
                file_ext = os.path.splitext(file_path)[1].lower()
                
                # Проверяем поддержку форматов
                if file_ext == '.pdf' and not PDF_SUPPORT:
                    messagebox.showerror("Ошибка", "Поддержка PDF не активирована. Установите PyPDF2: pip install PyPDF2")
                    continue
                elif file_ext in ['.doc', '.docx'] and not DOCX_SUPPORT:
                    messagebox.showerror("Ошибка", "Поддержка DOC/DOCX не активирована. Установите python-docx: pip install python-docx")
                    continue
                    
                self.loaded_files.append(file_path)
                filename = os.path.basename(file_path)
                self.files_listbox.insert(tk.END, filename)
        
        self.update_status(f"Загружено файлов: {len(self.loaded_files)}")
    
    def show_instructions(self):
        """Показать инструкцию по использованию"""
        instructions = """
KML Converter - Инструкция по использованию

1. ВЫБОР ФОРМАТА КООРДИНАТ
   - Стандартный: 564144N0523236E (градусы, минуты, секунды)
   - Географический: 56.1234, 52.5678 (десятичные градусы) или 56;331195 с. ш. ; 46;545811 в. д.

2. РУЧНОЙ ВВОД КООРДИНАТ
   - Вставляйте координаты в выбранном формате
   - Каждая координата должна быть на новой строке
   - Поддерживаемые форматы:
     * Отдельные точки: 524651N0500052E или 56.1234, 52.5678 или 56;331195 с. ш. ; 46;545811 в. д.
     * Круговые зоны: ОКРУЖНОСТЬ РАДИУС 10 КМ ЦЕНТР [координаты]
     * Полигоны: РАЙОН [координата1] [координата2] [координата3] ...

3. ЗАГРУЗКА ИЗ ФАЙЛОВ
   - Поддерживаются PDF, TXT, DOC и DOCX файлы
   - Программа автоматически найдет координаты в тексте файлов
   - Для PDF файлов требуется установленный PyPDF2
   - Для DOC/DOCX файлов требуется установленный python-docx

4. НАСТРОЙКА ЦВЕТОВ
   - Можно настроить цвета для разных типов объектов
   - Доступны стандартные цвета и выбор произвольного цвета

5. ОБРАБОТКА ДАННЫХ
   - Нажмите "Обработать данные" для анализа введенных координат
   - Результаты отобразятся в правой панели

6. ЭКСПОРТ KML
   - После обработки нажмите "Экспорт KML" для сохранения файла
   - KML файл можно открыть в Google Earth, Google Maps и других приложениях

ВАЖНО: 
   - Отдельные точки НЕ будут отображаться в KML файле
   - В KML файл экспортируются только круговые зоны и полигоны
   - Это сделано для улучшения восприятия представления

ПОДДЕРЖИВАЕМЫЕ ФОРМАТЫ КООРДИНАТ:
   СТАНДАРТНЫЙ:
   - 524651N0500052E (градусы, минуты, секунды с указанием направления)
   - Широта: 6 цифр + N/S (52°46'51"N)
   - Долгота: 7 цифр + E/W (050°00'52"E)

   ГЕОГРАФИЧЕСКИЙ:
   - 56.1234, 52.5678 (десятичные градусы)
   - 56;331195 с. ш. ; 46;545811 в. д. (градусы с дробной частью и указанием направлений)
   - Разделитель - запятая, точка с запятой или пробел
   - Можно использовать отрицательные значения для юга/запада

СОВЕТЫ:
   - Для копирования/вставки используйте Ctrl+C/Ctrl+V (работает в русской раскладке)
   - Можно одновременно обрабатывать данные из ручного ввода и файлов
   - Круговые зоны создаются с указанным радиусом в километрах
   - Полигоны должны содержать не менее 3 точек
        """
        
        # Создаем новое окно для инструкции
        instructions_window = tk.Toplevel(self.root)
        instructions_window.title("Инструкция по использованию KML Converter")
        instructions_window.geometry("600x500")
        instructions_window.resizable(True, True)
        
        # Создаем фрейм с прокруткой
        main_frame = ttk.Frame(instructions_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Добавляем текстовое поле с прокруткой
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        text_widget = tk.Text(text_frame, wrap=tk.WORD, padx=10, pady=10, font=('Arial', 10))
        scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Вставляем текст инструкции
        text_widget.insert(1.0, instructions)
        text_widget.config(state=tk.DISABLED)  # Делаем текст только для чтения
        
        # Кнопка закрытия
        close_button = ttk.Button(main_frame, text="Закрыть", 
                                 command=instructions_window.destroy)
        close_button.pack(pady=10)
    
    def process_data(self):
        """Обработка всех данных (ручной ввод + файлы)"""
        if not KML_SUPPORT:
            messagebox.showerror("Ошибка", "simplekml не установлен. Установите: pip install simplekml")
            return
        
        # Собираем все данные
        all_data = []
        
        # Обрабатываем ручной ввод
        manual_text = self.coord_text.get(1.0, tk.END).strip()
        if manual_text:
            manual_data = self.extract_data_from_text(manual_text)
            if manual_data['coordinates'] or manual_data['circles'] or manual_data['polygons']:
                all_data.append({
                    'source': 'manual',
                    'data': manual_data
                })
        
        # Обрабатываем файлы
        if self.loaded_files:
            for file_path in self.loaded_files:
                try:
                    file_data = self.parse_file(file_path)
                    if file_data:
                        all_data.append({
                            'source': f'file:{os.path.basename(file_path)}',
                            'data': file_data
                        })
                except Exception as e:
                    logging.error(f"Ошибка обработки файла {file_path}: {e}")
        
        if not all_data:
            messagebox.showwarning("Внимание", "Нет данных для обработки")
            return
        
        self.update_status("Обработка данных...")
        self.result_text.config(state=tk.NORMAL)
        self.result_text.delete(1.0, tk.END)
        
        try:
            # Объединяем все данные
            combined_data = {
                'coordinates': [],
                'circles': [],
                'polygons': []
            }
            
            manual_count = 0
            file_count = 0
            
            for item in all_data:
                data = item['data']
                combined_data['coordinates'].extend(data['coordinates'])
                combined_data['circles'].extend(data['circles'])
                combined_data['polygons'].extend(data['polygons'])
                
                if item['source'] == 'manual':
                    manual_count += 1
                elif item['source'].startswith('file:'):
                    file_count += 1
            
            self.processed_data = combined_data
            self.kml_data = self.create_kml_data(combined_data)
            
            # Показываем результаты
            self.result_text.insert(tk.END, f"Обработано источников: {len(all_data)}\n")
            self.result_text.insert(tk.END, f"- Ручной ввод: {manual_count}\n")
            self.result_text.insert(tk.END, f"- Файлов: {file_count}\n\n")
            
            self.result_text.insert(tk.END, f"Обнаружено:\n")
            self.result_text.insert(tk.END, f"- Координат: {len(combined_data['coordinates'])} (не отображаются в KML)\n")
            self.result_text.insert(tk.END, f"- Круговых зон: {len(combined_data['circles'])}\n")
            self.result_text.insert(tk.END, f"- Полигонов: {len(combined_data['polygons'])}\n\n")
            
            # Детали полигонов
            for i, polygon in enumerate(combined_data['polygons']):
                self.result_text.insert(tk.END, f"- Полигон {i+1}: {len(polygon)} точек\n")
            
            # Детали круговых зон
            for i, circle in enumerate(combined_data['circles']):
                self.result_text.insert(tk.END, f"- Круговая зона {i+1}: радиус {circle['radius_km']} км\n")
            
            self.result_text.insert(tk.END, "Можно экспортировать в KML\n")
            
            self.update_status("Обработка завершена")
            
        except Exception as e:
            error_msg = f"Ошибка обработки: {str(e)}"
            self.result_text.insert(tk.END, error_msg)
            self.update_status("Ошибка обработки")
            logging.error(error_msg)
        
        self.result_text.config(state=tk.DISABLED)
    
    def parse_file(self, file_path):
        """Парсинг файла (PDF, TXT, DOC, DOCX)"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                if not PDF_SUPPORT:
                    messagebox.showerror("Ошибка", "PyPDF2 не установлен. Установите: pip install PyPDF2")
                    return None
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                
                return self.extract_data_from_text(text)
                
            elif file_ext == '.txt':
                # Чтение TXT файла с различными кодировками
                encodings = ['utf-8', 'windows-1251', 'cp1251', 'iso-8859-1']
                text = ""
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        break
                    except UnicodeDecodeError:
                        continue
                
                return self.extract_data_from_text(text)
            
            elif file_ext in ['.doc', '.docx']:
                if not DOCX_SUPPORT:
                    messagebox.showerror("Ошибка", "python-docx не установлен. Установите: pip install python-docx")
                    return None
                
                return self.parse_docx_file(file_path)
                
            else:
                logging.warning(f"Неподдерживаемый формат файла: {file_path}")
                return None
                
        except Exception as e:
            logging.error(f"Ошибка парсинга файла {file_path}: {e}")
            return None
    
    def parse_docx_file(self, file_path):
        """Парсинг DOCX файла"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return self.extract_data_from_text(text)
        except Exception as e:
            logging.error(f"Ошибка парсинга DOCX файла {file_path}: {e}")
            return None
    
    def extract_data_from_text(self, text):
        """Извлечение данных из текста"""
        data = {
            'coordinates': [],
            'circles': [],
            'polygons': []
        }
        
        if self.coord_format == "standard":
            # УЛУЧШЕННЫЙ поиск координат - более гибкий паттерн
            coord_pattern = r'\b(\d{6}[NSСЮ]\d{7}[EWЗВ])\b'
            coord_matches = re.findall(coord_pattern, text, re.IGNORECASE)
            
            for coord in coord_matches:
                parsed_coord = self.parse_coordinate(coord)
                if parsed_coord:
                    data['coordinates'].append(parsed_coord)
            
            # УЛУЧШЕННЫЙ поиск круговых зон во взлет/посадке
            # Ищем все варианты фразы взлет/посадка
            takeoff_patterns = [
                r'Взлет/посадка[:\s]*([^\.\n]+?)(?:набор высоты|снижение|$|\.)',
                r'Взлет[:\s]*([^\.\n]+?)(?:набор высоты|снижение|$|\.)',
                r'посадка[:\s]*([^\.\n]+?)(?:набор высоты|снижение|$|\.)'
            ]
            
            for pattern in takeoff_patterns:
                takeoff_matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
                for takeoff_text in takeoff_matches:
                    # Ищем радиус в метрах
                    radius_match = re.search(r'радиусом?\s*(\d+)\s*м', takeoff_text, re.IGNORECASE)
                    if radius_match:
                        radius_m = int(radius_match.group(1))
                        radius_km = radius_m / 1000.0  # Конвертируем в километры
                        
                        # Ищем все координаты во взлет/посадка
                        takeoff_coords = re.findall(r'\b(\d{6}[NSСЮ]\d{7}[EWЗВ])\b', takeoff_text, re.IGNORECASE)
                        for coord in takeoff_coords:
                            center_coord = self.parse_coordinate(coord)
                            if center_coord:
                                data['circles'].append({
                                    'center': center_coord,
                                    'radius_km': radius_km
                                })
            
            # УЛУЧШЕННЫЙ поиск полигонов
            polygon_patterns = [
                r'Район полета[:\s]*([^\.\n]+?)(?:Высота|Время|$|\.)',
                r'Район полетов[:\s]*([^\.\n]+?)(?:Высота|Время|$|\.)',
                r'Район[:\s]*([^\.\n]+?)(?:Высота|Время|$|\.)'
            ]
            
            for pattern in polygon_patterns:
                polygon_matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
                for polygon_coords_text in polygon_matches:
                    # Ищем все координаты в тексте полигона
                    coords_list = re.findall(r'\b(\d{6}[NSСЮ]\d{7}[EWЗВ])\b', polygon_coords_text, re.IGNORECASE)
                    polygon_points = []
                    for coord in coords_list:
                        parsed_coord = self.parse_coordinate(coord)
                        if parsed_coord:
                            polygon_points.append(parsed_coord)
                    
                    if len(polygon_points) >= 3:
                        data['polygons'].append(polygon_points)
            
            # ДОПОЛНИТЕЛЬНЫЙ поиск: координаты в формате с буквой C вместо N и B вместо E
            alt_coord_pattern = r'\b(\d{6}[CС]\d{7}[BВ])\b'
            alt_coord_matches = re.findall(alt_coord_pattern, text, re.IGNORECASE)
            
            for coord in alt_coord_matches:
                # Заменяем C на N и B на E для парсинга
                normalized_coord = coord.replace('C', 'N').replace('С', 'N').replace('B', 'E').replace('В', 'E')
                parsed_coord = self.parse_coordinate(normalized_coord)
                if parsed_coord:
                    data['coordinates'].append(parsed_coord)
                    
        else:  # geographic format
            # Поиск отдельных координат (географический формат - десятичные градусы)
            coord_pattern = r'(-?\d+\.\d+)[,\s]+(-?\d+\.\d+)'
            coord_matches = re.findall(coord_pattern, text)
            
            for lat, lon in coord_matches:
                parsed_coord = self.parse_geographic_coordinate(lat, lon)
                if parsed_coord:
                    data['coordinates'].append(parsed_coord)
            
            # Поиск отдельных координат (географический формат - с точкой с запятой и указанием направлений)
            coord_pattern_semicolon = r'(\d+)[;](\d+)\s*с\.\s*ш\.\s*[;]\s*(\d+)[;](\d+)\s*в\.\s*д\.'
            coord_matches_semicolon = re.findall(coord_pattern_semicolon, text, re.IGNORECASE)
            
            for lat_deg, lat_frac, lon_deg, lon_frac in coord_matches_semicolon:
                parsed_coord = self.parse_semicolon_coordinate(lat_deg, lat_frac, lon_deg, lon_frac)
                if parsed_coord:
                    data['coordinates'].append(parsed_coord)
            
            # Поиск круговых зон (географический формат - десятичные градусы)
            circle_pattern = r'ОКРУЖНОСТЬ РАДИУС\s+(\d+)\s+КМ ЦЕНТР\s+(-?\d+\.\d+)[,\s]+(-?\d+\.\d+)'
            circle_matches = re.findall(circle_pattern, text)
            
            for radius, lat, lon in circle_matches:
                center_coord = self.parse_geographic_coordinate(lat, lon)
                if center_coord:
                    data['circles'].append({
                        'center': center_coord,
                        'radius_km': int(radius)
                    })
            
            # Поиск круговых зон (географический формат - с точкой с запятой)
            circle_pattern_semicolon = r'ОКРУЖНОСТЬ РАДИУС\s+(\d+)\s+КМ ЦЕНТР\s+(\d+)[;](\d+)\s*с\.\s*ш\.\s*[;]\s*(\d+)[;](\d+)\s*в\.\s*д\.'
            circle_matches_semicolon = re.findall(circle_pattern_semicolon, text, re.IGNORECASE)
            
            for radius, lat_deg, lat_frac, lon_deg, lon_frac in circle_matches_semicolon:
                center_coord = self.parse_semicolon_coordinate(lat_deg, lat_frac, lon_deg, lon_frac)
                if center_coord:
                    data['circles'].append({
                        'center': center_coord,
                        'radius_km': int(radius)
                    })
            
            # Поиск полигонов (географический формат - десятичные градусы)
            polygon_pattern = r'РАЙОН\s+((?:(?:-?\d+\.\d+)[,\s]+(?:-?\d+\.\d+)\s*)+)'
            polygon_matches = re.findall(polygon_pattern, text)
            
            for polygon_coords in polygon_matches:
                coords_list = re.findall(r'(-?\d+\.\d+)[,\s]+(-?\d+\.\d+)', polygon_coords)
                polygon_points = []
                for lat, lon in coords_list:
                    parsed_coord = self.parse_geographic_coordinate(lat, lon)
                    if parsed_coord:
                        polygon_points.append(parsed_coord)
                
                if len(polygon_points) >= 3:
                    data['polygons'].append(polygon_points)
            
            # Поиск полигонов (географический формат - с точкой с запятой)
            polygon_pattern_semicolon = r'РАЙОН\s+((?:(?:\d+)[;](\d+)\s*с\.\s*ш\.\s*[;]\s*(?:\d+)[;](\d+)\s*в\.\s*д\.\s*)+)'
            polygon_matches_semicolon = re.findall(polygon_pattern_semicolon, text, re.IGNORECASE)
            
            for polygon_coords in polygon_matches_semicolon:
                coords_list = re.findall(r'(\d+)[;](\d+)\s*с\.\s*ш\.\s*[;]\s*(\d+)[;](\d+)\s*в\.\s*д\.', polygon_coords)
                polygon_points = []
                for lat_deg, lat_frac, lon_deg, lon_frac in coords_list:
                    parsed_coord = self.parse_semicolon_coordinate(lat_deg, lat_frac, lon_deg, lon_frac)
                    if parsed_coord:
                        polygon_points.append(parsed_coord)
                
                if len(polygon_points) >= 3:
                    data['polygons'].append(polygon_points)
        
        return data
    
    def parse_coordinate(self, coord_str):
        """Парсинг координат из стандартного строкового формата"""
        try:
            # Формат: 524651N0500052E или 612000С0765645В (с русскими буквами)
            # Нормализуем русские буквы к английским
            normalized_coord = coord_str.replace('С', 'N').replace('Ю', 'S').replace('В', 'E').replace('З', 'W')
            
            lat_match = re.search(r'(\d{2})(\d{2})(\d{2})([NS])', normalized_coord)
            lon_match = re.search(r'(\d{3})(\d{2})(\d{2})([EW])', normalized_coord)
            
            if lat_match and lon_match:
                lat_deg = int(lat_match.group(1))
                lat_min = int(lat_match.group(2))
                lat_sec = int(lat_match.group(3))
                lat_dir = lat_match.group(4)
                
                lon_deg = int(lon_match.group(1))
                lon_min = int(lon_match.group(2))
                lon_sec = int(lon_match.group(3))
                lon_dir = lon_match.group(4)
                
                # Преобразование в десятичные градусы
                lat_decimal = lat_deg + lat_min/60 + lat_sec/3600
                lon_decimal = lon_deg + lon_min/60 + lon_sec/3600
                
                if lat_dir == 'S':
                    lat_decimal = -lat_decimal
                if lon_dir == 'W':
                    lon_decimal = -lon_decimal
                
                return {
                    'original': coord_str,
                    'decimal': (lat_decimal, lon_decimal),
                    'degrees_minutes_seconds': {
                        'lat': f"{lat_deg}°{lat_min:02d}'{lat_sec:02d}\"{lat_dir}",
                        'lon': f"{lon_deg}°{lon_min:02d}'{lon_sec:02d}\"{lon_dir}"
                    }
                }
        except Exception as e:
            logging.error(f"Ошибка парсинга координаты {coord_str}: {e}")
        
        return None
    
    def parse_geographic_coordinate(self, lat_str, lon_str):
        """Парсинг координат из географического формата (десятичные градусы)"""
        try:
            lat_decimal = float(lat_str)
            lon_decimal = float(lon_str)
            
            # Форматируем оригинальную строку
            original = f"{lat_decimal:.6f}, {lon_decimal:.6f}"
            
            return {
                'original': original,
                'decimal': (lat_decimal, lon_decimal),
                'degrees_minutes_seconds': {
                    'lat': f"{abs(lat_decimal):.6f}° {'N' if lat_decimal >= 0 else 'S'}",
                    'lon': f"{abs(lon_decimal):.6f}° {'E' if lon_decimal >= 0 else 'W'}"
                }
            }
        except Exception as e:
            logging.error(f"Ошибка парсинга географической координаты {lat_str}, {lon_str}: {e}")
        
        return None
    
    def parse_semicolon_coordinate(self, lat_deg, lat_frac, lon_deg, lon_frac):
        """Парсинг координат из формата с точкой с запятой и указанием направлений"""
        try:
            # Формат: 56;331195 с. ш. ; 46;545811 в. д.
            # Преобразуем в десятичные градусы: градусы + дробная часть / 1000000
            lat_decimal = int(lat_deg) + int(lat_frac) / 1000000
            lon_decimal = int(lon_deg) + int(lon_frac) / 1000000
            
            # Форматируем оригинальную строку
            original = f"{lat_deg};{lat_frac} с. ш. ; {lon_deg};{lon_frac} в. д."
            
            return {
                'original': original,
                'decimal': (lat_decimal, lon_decimal),
                'degrees_minutes_seconds': {
                    'lat': f"{lat_deg}°{int(lat_frac)/10000:.4f}' N",
                    'lon': f"{lon_deg}°{int(lon_frac)/10000:.4f}' E"
                }
            }
        except Exception as e:
            logging.error(f"Ошибка парсинга координаты с точкой с запятой {lat_deg};{lat_frac}, {lon_deg};{lon_frac}: {e}")
        
        return None
    
    def create_kml_data(self, data):
        """Создание KML данных - только полигоны и круговые зоны"""
        kml = simplekml.Kml()
        
        # Добавляем круговые зоны
        for i, circle in enumerate(data['circles']):
            circle_points = self.create_circle_points(
                circle['center']['decimal'][0],  # lat
                circle['center']['decimal'][1],  # lon
                circle['radius_km']  # радиус
            )
            pol = kml.newpolygon(
                name=f"Круговая зона {i+1} (R={circle['radius_km']:.3f}км)",
                outerboundaryis=circle_points
            )
            circle_color = self.color_settings.get("circle_color", "blue")
            pol.style.polystyle.color = self.get_kml_color(circle_color, 60)
            pol.style.linestyle.color = self.get_kml_color(circle_color)
            pol.style.linestyle.width = 2
            pol.description = (
                f"<![CDATA["
                f"<b>Круговая зона {i+1}</b><br/>"
                f"<b>Центр:</b> {circle['center']['original']}<br/>"
                f"<b>Радиус:</b> {circle['radius_km']:.3f} км<br/>"
                f"]]>"
            )
        
        # Добавляем полигоны
        for i, polygon in enumerate(data['polygons']):
            # Правильный порядок для KML полигонов: (longitude, latitude)
            poly_coords = [(p['decimal'][1], p['decimal'][0]) for p in polygon]
            pol = kml.newpolygon(
                name=f"Полигон {i+1}",
                outerboundaryis=poly_coords
            )
            polygon_color = self.color_settings.get("polygon_color", "yellow")
            pol.style.polystyle.color = self.get_kml_color(polygon_color, 80)
            pol.style.linestyle.color = self.get_kml_color(polygon_color)
            pol.style.linestyle.width = 3
            pol.description = (
                f"<![CDATA["
                f"<b>Полигон {i+1}</b><br/>"
                f"<b>Точек в полигоне:</b> {len(polygon)}<br/>"
                f"]]>"
            )
        
        return kml
    
    def get_kml_color(self, color, alpha=255):
        """Преобразование цвета в формат KML (AABBGGRR)"""
        # Простые цвета
        color_map = {
            "red": simplekml.Color.red,
            "green": simplekml.Color.green,
            "blue": simplekml.Color.blue,
            "yellow": simplekml.Color.yellow,
            "orange": simplekml.Color.orange,
            "purple": simplekml.Color.purple,
            "pink": simplekml.Color.pink,
            "brown": simplekml.Color.brown,
            "black": simplekml.Color.black,
            "white": simplekml.Color.white,
            "cyan": simplekml.Color.cyan,
            "lime": simplekml.Color.lime
        }
        
        if color in color_map:
            return simplekml.Color.changealphaint(alpha, color_map[color])
        else:
            # Для пользовательских цветов в формате #RRGGBB
            try:
                if color.startswith('#'):
                    # Конвертируем #RRGGBB в AABBGGRR
                    rr = color[1:3]
                    gg = color[3:5]
                    bb = color[5:7]
                    aa = f"{alpha:02x}"
                    return f"{aa}{bb}{gg}{rr}"
            except:
                pass
        
        # По умолчанию синий
        return simplekml.Color.changealphaint(alpha, simplekml.Color.blue)
    
    def create_circle_points(self, lat, lon, radius_km, points=36):
        """Создание точек для круговой зоны"""
        coords = []
        R = 6371.0  # Радиус Земли в км
        
        for i in range(points + 1):  # +1 для замыкания круга
            angle = 2.0 * math.pi * i / points
            
            # Вычисление новой точки с учетом сферической геометрии
            lat_rad = math.radians(lat)
            lon_rad = math.radians(lon)
            
            new_lat = math.asin(math.sin(lat_rad) * math.cos(radius_km/R) + 
                               math.cos(lat_rad) * math.sin(radius_km/R) * math.cos(angle))
            new_lon = lon_rad + math.atan2(math.sin(angle) * math.sin(radius_km/R) * math.cos(lat_rad),
                                         math.cos(radius_km/R) - math.sin(lat_rad) * math.sin(new_lat))
            
            new_lat_deg = math.degrees(new_lat)
            new_lon_deg = math.degrees(new_lon)
            
            # Правильный порядок для KML: (longitude, latitude)
            coords.append((new_lon_deg, new_lat_deg))
        
        return coords
    
    def export_kml(self):
        """Экспорт KML файла"""
        if not self.kml_data:
            messagebox.showwarning("Внимание", "Сначала обработайте данные")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Сохранить KML файл",
            defaultextension=".kml",
            filetypes=[("KML files", "*.kml"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                # Добавляем метаданные
                self.kml_data.document.name = "KML Converter 1.4 - Созданные объекты"
                self.kml_data.document.description = "Сгенерировано KML Converter плагином"
                
                # Сохраняем файл
                self.kml_data.save(file_path)
                
                self.update_status(f"KML файл сохранен: {file_path}")
                
                # Показываем информацию
                messagebox.showinfo(
                    "Успех", 
                    f"KML файл успешно сохранен:\n{file_path}\n\n"
                    f"Создано объектов:\n"
                    f"• Круговых зон: {len(self.processed_data['circles'])}\n"
                    f"• Полигонов: {len(self.processed_data['polygons'])}\n\n"
                    f"Отдельные точки не экспортируются для улучшения восприятия."
                )
                
                # Показываем в логе
                self.result_text.config(state=tk.NORMAL)
                self.result_text.insert(tk.END, f"\n\nKML файл экспортирован: {file_path}")
                self.result_text.see(tk.END)
                self.result_text.config(state=tk.DISABLED)
                
            except Exception as e:
                error_msg = f"Ошибка сохранения KML: {str(e)}"
                messagebox.showerror("Ошибка", error_msg)
                logging.error(error_msg)
    
    def update_status(self, message):
        """Обновление статусной строки"""
        self.status_var.set(message)
        self.root.update_idletasks()
    
    def load_settings(self):
        """Загрузка настроек из файла"""
        try:
            settings_file = "kml_settings.json"
            if os.path.exists(settings_file):
                with open(settings_file, 'r', encoding='utf-8') as f:
                    loaded_settings = json.load(f)
                
                # Загружаем настройки цветов
                if "color_settings" in loaded_settings:
                    self.color_settings.update(loaded_settings["color_settings"])
                
                # Загружаем формат координат
                if "coord_format" in loaded_settings:
                    self.coord_format = loaded_settings["coord_format"]
                    
        except Exception as e:
            logging.error(f"Ошибка загрузки настроек: {e}")
            # В случае ошибки используем настройки по умолчанию
            self.save_settings()
    
    def save_settings(self):
        """Сохранение настроек в файл"""
        try:
            settings_file = "kml_settings.json"
            settings = {
                "color_settings": self.color_settings,
                "coord_format": self.coord_format
            }
            with open(settings_file, 'w', encoding='utf-8') as f:
                json.dump(settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logging.error(f"Ошибка сохранения настроек: {e}")

def get_plugin_class():
    return KMLCreatorPlugin